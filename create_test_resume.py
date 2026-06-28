import docx

def create_resume():
    doc = docx.Document()
    
    # Title
    doc.add_heading('Sushant Kumar', 0)
    
    # Contact Info
    doc.add_paragraph('Title: Senior Full Stack Engineer\nEmail: sushant.kumar@example.com | Phone: +91 98765 43210\nLinkedIn: linkedin.com/in/sushant | GitHub: github.com/sushant | Portfolio: sushant.dev')
    
    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Highly skilled Senior Full Stack Engineer with over 7 years of experience specializing in React, Node.js, Python, Supabase, and cloud architectures. Passionate about building performant, secure, and user-centric web applications.')
    
    # Experience
    doc.add_heading('Work Experience', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('Senior Software Engineer - Google').bold = True
    p1.add_run('\nBangalore, India | 2022-03 to Present\n')
    doc.add_paragraph('- Led the development of high-performance web applications using React and Node.js.\n- Optimized database queries, reducing response times by 35%.\n- Mentored junior engineers and implemented CI/CD pipelines.')
    
    p2 = doc.add_paragraph()
    p2.add_run('Full Stack Developer - Infosys').bold = True
    p2.add_run('\nPune, India | 2019-06 to 2022-02\n')
    doc.add_paragraph('- Developed features for enterprise client dashboards.\n- Built RESTful APIs using Express and integrated third-party payment gateways.\n- Reduced security vulnerabilities by 40% through code audits.')
    
    # Education
    doc.add_heading('Education', level=1)
    p3 = doc.add_paragraph()
    p3.add_run('B.Tech in Computer Science - IIT Delhi').bold = True
    p3.add_run('\nNew Delhi, India | 2015-07 to 2019-05\n')
    doc.add_paragraph('Graduated with Honors. Specialized in algorithms and database systems.')
    
    # Skills
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('React, Node.js, Python, Supabase, PostgreSQL, Docker, Git, REST APIs, AWS')
    
    # Projects
    doc.add_heading('Projects', level=1)
    p4 = doc.add_paragraph()
    p4.add_run('AI Resume Builder - Lead Architect').bold = True
    p4.add_run('\nhttps://ai-resume-builder.dev\n')
    doc.add_paragraph('Designed and built an intelligent resume generator utilizing Gemini to parse, scan, and optimize user resumes, generating real-time ATS critiques.')
    
    doc.save('my-resume.docx')
    print("Test resume docx created successfully!")

if __name__ == '__main__':
    create_resume()
